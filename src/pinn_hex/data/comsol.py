from __future__ import annotations

from dataclasses import asdict, dataclass
from pathlib import Path
import json

import numpy as np
import pandas as pd
from docx import Document

from pinn_hex.data.preprocessing import build_preprocessed_supervision, save_preprocessed_supervision
from pinn_hex.physics.double_pipe import Geometry

CSV_COLUMNS = ["x", "y", "z", "T"]


@dataclass
class DatasetSummary:
    name: str
    n_rows: int
    n_cols: int
    temperature_min_K: float
    temperature_max_K: float
    temperature_std_K: float
    x_min_m: float
    x_max_m: float
    y_min_m: float
    y_max_m: float
    z_min_m: float
    z_max_m: float
    duplicate_rows: int
    constant_temperature: bool
    metadata: dict


def read_temperature_csv(path: str | Path) -> pd.DataFrame:
    data = pd.read_csv(path, comment="%", header=None, names=CSV_COLUMNS)
    data["r"] = np.sqrt(data["x"] ** 2 + data["y"] ** 2)
    return data


def read_csv_metadata(path: str | Path) -> dict[str, str]:
    metadata: dict[str, str] = {}
    with Path(path).open("r", encoding="utf-8") as handle:
        for raw_line in handle:
            if not raw_line.startswith("%"):
                break
            line = raw_line[1:].strip()
            if "," not in line:
                continue
            key, value = line.split(",", 1)
            metadata[key.strip()] = value.strip().strip('"')
    return metadata


def summarize_temperature_csv(path: str | Path) -> DatasetSummary:
    data = read_temperature_csv(path)
    return DatasetSummary(
        name=Path(path).name,
        n_rows=int(data.shape[0]),
        n_cols=int(data.shape[1]),
        temperature_min_K=float(data["T"].min()),
        temperature_max_K=float(data["T"].max()),
        temperature_std_K=float(data["T"].std()),
        x_min_m=float(data["x"].min()),
        x_max_m=float(data["x"].max()),
        y_min_m=float(data["y"].min()),
        y_max_m=float(data["y"].max()),
        z_min_m=float(data["z"].min()),
        z_max_m=float(data["z"].max()),
        duplicate_rows=int(data[CSV_COLUMNS].duplicated().sum()),
        constant_temperature=bool(np.isclose(data["T"].std(), 0.0, atol=1e-9)),
        metadata=read_csv_metadata(path),
    )


def frames_match(left: pd.DataFrame, right: pd.DataFrame, atol: float = 1e-9) -> bool:
    if left.shape != right.shape:
        return False
    return bool(np.allclose(left[CSV_COLUMNS].to_numpy(), right[CSV_COLUMNS].to_numpy(), atol=atol, rtol=0.0))


def classify_regions(data: pd.DataFrame, geometry: Geometry) -> pd.DataFrame:
    result = data.copy()
    hot_limit = geometry.hot_radius_m + geometry.radial_tolerance_m
    cold_limit = geometry.annulus_outer_radius_m + geometry.radial_tolerance_m
    result["region"] = "outside"
    hot_mask = result["r"] <= hot_limit
    cold_mask = (
        (result["r"] > hot_limit)
        & (result["r"] <= cold_limit)
        & (result["z"].abs() <= geometry.cold_half_length_m + geometry.radial_tolerance_m)
    )
    result.loc[hot_mask, "region"] = "hot"
    result.loc[cold_mask, "region"] = "cold"
    return result


def _bin_profile(data: pd.DataFrame, z_min: float, z_max: float, bins: int, flow: str) -> pd.DataFrame:
    edges = np.linspace(z_min, z_max, bins + 1)
    labels = np.arange(bins)
    work = data.copy()
    work["bin"] = pd.cut(work["z"], bins=edges, labels=labels, include_lowest=True)
    profile = (
        work.dropna(subset=["bin"])
        .groupby("bin", observed=False)
        .agg(
            z_mean=("z", "mean"),
            z_min=("z", "min"),
            z_max=("z", "max"),
            T_mean=("T", "mean"),
            T_std=("T", "std"),
            n=("T", "size"),
        )
        .reset_index(drop=True)
    )
    if flow == "hot":
        profile["s_flow"] = (z_max - profile["z_mean"]) / (z_max - z_min)
    elif flow == "cold":
        profile["s_flow"] = (profile["z_mean"] - z_min) / (z_max - z_min)
    else:
        raise ValueError(f"Unknown flow label: {flow}")
    return profile.sort_values("s_flow").reset_index(drop=True)


def face_statistics(data: pd.DataFrame, geometry: Geometry) -> dict[str, dict[str, float | int]]:
    stats: dict[str, dict[str, float | int]] = {}
    faces = {
        "hot_inlet": geometry.hot_half_length_m,
        "hot_outlet": -geometry.hot_half_length_m,
        "cold_outlet_side": geometry.cold_half_length_m,
        "cold_inlet_side": -geometry.cold_half_length_m,
    }
    hot_limit = geometry.hot_radius_m + geometry.radial_tolerance_m
    cold_limit = geometry.annulus_outer_radius_m + geometry.radial_tolerance_m
    for name, z_target in faces.items():
        window = data[np.isclose(data["z"], z_target, atol=1e-4)].copy()
        if window.empty:
            continue
        if "hot" in name:
            subset = window[window["r"] <= hot_limit]
        else:
            subset = window[(window["r"] > hot_limit) & (window["r"] <= cold_limit)]
        if subset.empty:
            continue
        stats[name] = {
            "n": int(len(subset)),
            "T_mean_K": float(subset["T"].mean()),
            "T_min_K": float(subset["T"].min()),
            "T_max_K": float(subset["T"].max()),
        }
    return stats


def infer_geometry(solution_data: pd.DataFrame) -> dict[str, float]:
    hot_slice = solution_data[np.isclose(solution_data["z"], solution_data["z"].max(), atol=1e-6)].copy()
    cold_slice = solution_data[np.isclose(solution_data["z"], 0.165, atol=1e-6)].copy()
    return {
        "hot_half_length_m": float(solution_data["z"].max()),
        "hot_radius_m": float(hot_slice["r"].max()) if not hot_slice.empty else np.nan,
        "cold_half_length_m": 0.165,
        "cold_outer_radius_m": float(cold_slice["r"].max()) if not cold_slice.empty else np.nan,
    }


def extract_reference_summary(path: str | Path) -> dict:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    keywords = [
        "Heat Transfer in Fluids",
        "Turbulent Flow, k-ω",
        "Upstream temperature",
        "Normal inflow velocity",
        "Initial temperature",
        "Mesh vertices",
        "Heat transport turbulence model",
    ]
    matching_paragraphs = [text for text in paragraphs if any(key.lower() in text.lower() for key in keywords)]
    matching_tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        block = " ".join(" | ".join(row) for row in rows)
        if any(key.lower() in block.lower() for key in keywords):
            matching_tables.append(rows[:12])
    return {
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "matching_paragraphs": matching_paragraphs[:60],
        "matching_tables": matching_tables[:20],
    }


def build_case_artifacts(config: dict) -> dict:
    solution_path = Path(config["paths"]["solution_csv"])
    mesh_path = Path(config["paths"]["mesh_csv"])
    duplicate_path = Path(config["paths"]["duplicate_csv"])
    reference_path = Path(config["paths"]["reference_docx"])
    geometry = Geometry(
        hot_half_length_m=float(config["geometry"]["hot_half_length_m"]),
        cold_half_length_m=float(config["geometry"]["cold_half_length_m"]),
        hot_radius_m=float(config["geometry"]["hot_radius_m"]),
        annulus_outer_radius_m=float(config["geometry"]["annulus_outer_radius_m"]),
        radial_tolerance_m=float(config["geometry"]["radial_tolerance_m"]),
    )
    solution = classify_regions(read_temperature_csv(solution_path), geometry)
    mesh = classify_regions(read_temperature_csv(mesh_path), geometry)
    duplicate = read_temperature_csv(duplicate_path)
    hot_profile = _bin_profile(
        solution[solution["region"] == "hot"],
        -geometry.hot_half_length_m,
        geometry.hot_half_length_m,
        int(config["preprocessing"]["hot_bins"]),
        flow="hot",
    )
    cold_profile = _bin_profile(
        solution[solution["region"] == "cold"],
        -geometry.cold_half_length_m,
        geometry.cold_half_length_m,
        int(config["preprocessing"]["cold_bins"]),
        flow="cold",
    )
    from pinn_hex.physics.double_pipe import OperatingPoint

    operating = OperatingPoint(
        hot_inlet_temperature_K=float(config["reference_conditions"]["hot_inlet_temperature_K"]),
        cold_inlet_temperature_K=float(config["reference_conditions"]["cold_inlet_temperature_K"]),
        initial_temperature_K=float(config["reference_conditions"]["initial_temperature_K"]),
        inlet_velocity_hot_m_per_s=float(config["reference_conditions"]["inlet_velocity_hot_m_per_s"]),
        inlet_velocity_cold_m_per_s=float(config["reference_conditions"]["inlet_velocity_cold_m_per_s"]),
        density_kg_per_m3=float(config["reference_conditions"]["density_kg_per_m3"]),
        cp_J_per_kgK=float(config["reference_conditions"]["cp_J_per_kgK"]),
    )
    preprocessed = build_preprocessed_supervision(
        hot_profile=hot_profile,
        cold_profile=cold_profile,
        geometry=geometry,
        operating=operating,
        cfg=config["preprocessing"],
    )
    return {
        "solution": solution,
        "mesh": mesh,
        "hot_profile": hot_profile,
        "cold_profile": cold_profile,
        "preprocessed": preprocessed,
        "summary": {
            "solution_csv": asdict(summarize_temperature_csv(solution_path)),
            "mesh_csv": asdict(summarize_temperature_csv(mesh_path)),
            "duplicate_csv": asdict(summarize_temperature_csv(duplicate_path)),
            "duplicate_matches_solution": frames_match(solution[CSV_COLUMNS], duplicate[CSV_COLUMNS]),
            "inferred_geometry": infer_geometry(solution),
            "face_statistics": face_statistics(solution, geometry),
            "region_counts": solution["region"].value_counts().to_dict(),
            "mesh_region_counts": mesh["region"].value_counts().to_dict(),
            "reference_summary": extract_reference_summary(reference_path),
            "preprocessing_summary": preprocessed["summary"],
        },
    }


def save_case_artifacts(artifacts: dict, processed_dir: str | Path) -> None:
    output_dir = Path(processed_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    artifacts["hot_profile"].to_csv(output_dir / "hot_profile.csv", index=False)
    artifacts["cold_profile"].to_csv(output_dir / "cold_profile.csv", index=False)
    artifacts["solution"].to_csv(output_dir / "solution_classified.csv", index=False)
    artifacts["mesh"].to_csv(output_dir / "mesh_classified.csv", index=False)
    save_preprocessed_supervision(artifacts["preprocessed"], output_dir)
    with (output_dir / "analysis_summary.json").open("w", encoding="utf-8") as handle:
        json.dump(artifacts["summary"], handle, indent=2)
